# -*- coding: utf-8 -*-
"""
Script tạo file .docx mẫu cho Tendoo documentation
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_tendoo_document():
    """Tạo tài liệu mẫu theo cấu trúc Tendoo"""

    doc = Document()

    # Tiêu đề chính
    title = doc.add_paragraph("TÀI LIỆU HƯỚNG DẪN SỬ DỤNG TENDOO APP")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    # Phần mở đầu
    doc.add_paragraph("Phiên bản: 1.0")
    doc.add_paragraph("Ngày phát hành: 15/12/2025")
    doc.add_paragraph()
    doc.add_paragraph(
        "Tài liệu này hướng dẫn chi tiết cách sử dụng Tendoo App - "
        "Giải pháp quản lý cửa hàng và bán hàng toàn diện."
    )
    doc.add_paragraph()

    # 1. Cài đặt cửa hàng
    doc.add_paragraph("1. Cài đặt cửa hàng")
    doc.add_paragraph(
        "Phần này hướng dẫn bạn thiết lập các thông tin cơ bản cho cửa hàng của mình."
    )

    # 1.1. Cửa hàng
    doc.add_paragraph("1.1. Cửa hàng")
    doc.add_paragraph(
        "Quản lý thông tin chung về cửa hàng, bao gồm thông tin cửa hàng, "
        "website bán hàng, phương thức thanh toán và đăng ký nhà bán."
    )

    # 1.1.1. Thông tin cửa hàng
    doc.add_paragraph("1.1.1. Thông tin cửa hàng")
    doc.add_paragraph("Để cập nhật thông tin cửa hàng:")
    doc.add_paragraph("1 Vào menu Cài đặt > Cửa hàng > Thông tin cửa hàng")
    doc.add_paragraph("2 Điền các thông tin sau:")
    doc.add_paragraph("- Tên cửa hàng")
    doc.add_paragraph("- Địa chỉ")
    doc.add_paragraph("- Số điện thoại")
    doc.add_paragraph("- Email liên hệ")
    doc.add_paragraph("3 Nhấn nút Lưu để hoàn tất")
    doc.add_paragraph()
    doc.add_paragraph("Lưu ý:")
    doc.add_paragraph("+ Tên cửa hàng sẽ hiển thị trên hóa đơn")
    doc.add_paragraph("+ Email sẽ được sử dụng để nhận thông báo")

    # 1.1.2. Website bán hàng
    doc.add_paragraph("1.1.2. Website bán hàng")
    doc.add_paragraph(
        "Tendoo cho phép bạn tạo website bán hàng trực tuyến một cách dễ dàng. "
        "Hệ thống sẽ tự động đồng bộ sản phẩm và đơn hàng giữa app và website."
    )
    doc.add_paragraph("Các bước thiết lập:")
    doc.add_paragraph("1 Chọn mẫu giao diện phù hợp")
    doc.add_paragraph("2 Tùy chỉnh màu sắc và logo")
    doc.add_paragraph("3 Cấu hình tên miền")
    doc.add_paragraph("- Sử dụng tên miền con miễn phí: shop.tendoo.vn")
    doc.add_paragraph("- Hoặc kết nối tên miền riêng")
    doc.add_paragraph("4 Kích hoạt website")

    # 1.1.3. Phương thức thanh toán
    doc.add_paragraph("1.1.3. Phương thức thanh toán")
    doc.add_paragraph("Tendoo hỗ trợ nhiều phương thức thanh toán:")
    doc.add_paragraph("- Tiền mặt")
    doc.add_paragraph("- Chuyển khoản ngân hàng")
    doc.add_paragraph("- Ví điện tử (MoMo, ZaloPay, VNPay)")
    doc.add_paragraph("- Thẻ tín dụng/ghi nợ")
    doc.add_paragraph()
    doc.add_paragraph("Để cấu hình:")
    doc.add_paragraph("1 Vào Cài đặt > Thanh toán")
    doc.add_paragraph("2 Chọn phương thức muốn kích hoạt")
    doc.add_paragraph("3 Điền thông tin tài khoản")
    doc.add_paragraph("+ Số tài khoản ngân hàng")
    doc.add_paragraph("+ API key của cổng thanh toán")

    # 1.1.4. Đăng ký nhà bán
    doc.add_paragraph("1.1.4. Đăng ký nhà bán")
    doc.add_paragraph(
        "Nếu bạn muốn bán hàng trên sàn thương mại điện tử Tendoo Marketplace, "
        "cần đăng ký trở thành nhà bán chính thức."
    )
    doc.add_paragraph("Quy trình đăng ký:")
    doc.add_paragraph("1 Điền form đăng ký nhà bán")
    doc.add_paragraph("2 Cung cấp giấy tờ pháp lý")
    doc.add_paragraph("- Giấy phép kinh doanh")
    doc.add_paragraph("- CMND/CCCD chủ cửa hàng")
    doc.add_paragraph("3 Chờ phê duyệt (1-3 ngày làm việc)")
    doc.add_paragraph("4 Nhận thông báo kết quả qua email")

    # 1.2. Tối ưu bán hàng
    doc.add_paragraph("1.2. Tối ưu bán hàng")
    doc.add_paragraph(
        "Phần này giúp bạn thiết lập quy trình bán hàng hiệu quả, "
        "quản lý thông tin sản phẩm và tùy chỉnh mẫu hóa đơn."
    )

    # 1.2.1. Quy trình bán hàng
    doc.add_paragraph("1.2.1. Quy trình bán hàng")
    doc.add_paragraph(
        "Tendoo hỗ trợ 2 quy trình bán hàng chính tùy theo loại hình kinh doanh của bạn."
    )

    # 1.2.1.1. Quy trình cho FnB
    doc.add_paragraph("1.2.1.1. Quy trình bán hàng cho chủ shop FnB")
    doc.add_paragraph(
        "Dành cho các cửa hàng thức ăn, đồ uống với đặc thù phục vụ nhanh."
    )
    doc.add_paragraph("Các bước:")
    doc.add_paragraph("1 Khách hàng order tại quầy")
    doc.add_paragraph("2 Nhân viên nhập đơn trên app")
    doc.add_paragraph("- Chọn sản phẩm")
    doc.add_paragraph("- Chọn size, topping")
    doc.add_paragraph("+ Thêm ghi chú nếu cần")
    doc.add_paragraph("3 Gửi order vào bếp")
    doc.add_paragraph("4 Bếp chuẩn bị món")
    doc.add_paragraph("5 Nhân viên phục vụ mang món cho khách")
    doc.add_paragraph("6 Thu tiền và in hóa đơn")

    # 1.2.1.2. Quy trình cho bán lẻ
    doc.add_paragraph("1.2.1.2. Quy trình bán hàng cho chủ shop bán lẻ")
    doc.add_paragraph("Dành cho các cửa hàng bán lẻ, thời trang, điện tử, v.v.")
    doc.add_paragraph("Các bước:")
    doc.add_paragraph("1 Khách hàng chọn sản phẩm")
    doc.add_paragraph("2 Nhân viên quét mã vạch hoặc chọn sản phẩm")
    doc.add_paragraph("3 Kiểm tra tồn kho")
    doc.add_paragraph("- Nếu hết hàng, đề xuất sản phẩm thay thế")
    doc.add_paragraph("4 Áp dụng khuyến mãi (nếu có)")
    doc.add_paragraph("5 Thu tiền")
    doc.add_paragraph("6 In hóa đơn và giao hàng")

    # 1.2.2. Thông tin sản phẩm
    doc.add_paragraph("1.2.2. Thông tin sản phẩm")
    doc.add_paragraph("Quản lý thông tin sản phẩm đầy đủ giúp tăng doanh số.")
    doc.add_paragraph("Các thông tin cần có:")
    doc.add_paragraph("1 Tên sản phẩm (rõ ràng, dễ tìm)")
    doc.add_paragraph("2 Mã SKU (mã định danh duy nhất)")
    doc.add_paragraph("3 Giá bán")
    doc.add_paragraph("- Giá gốc")
    doc.add_paragraph("- Giá khuyến mãi")
    doc.add_paragraph("4 Hình ảnh sản phẩm")
    doc.add_paragraph("5 Mô tả chi tiết")
    doc.add_paragraph("6 Danh mục sản phẩm")

    # 1.2.3. Mẫu hóa đơn
    doc.add_paragraph("1.2.3. Mẫu hóa đơn")
    doc.add_paragraph("Tendoo cung cấp 4 mẫu hóa đơn chuyên nghiệp.")

    # 1.2.3.1
    doc.add_paragraph("1.2.3.1. Mẫu hóa đơn 1: Khổ 80mm – mẫu mặc định của hệ thống")
    doc.add_paragraph("Đây là mẫu hóa đơn cơ bản, phù hợp với hầu hết các loại hình kinh doanh.")
    doc.add_paragraph("Đặc điểm:")
    doc.add_paragraph("- Khổ giấy: 80mm")
    doc.add_paragraph("- Bố cục đơn giản, dễ đọc")
    doc.add_paragraph("- Hiển thị đầy đủ thông tin:")
    doc.add_paragraph("+ Logo và tên cửa hàng")
    doc.add_paragraph("+ Danh sách sản phẩm")
    doc.add_paragraph("+ Tổng tiền và phương thức thanh toán")

    # 1.2.3.2
    doc.add_paragraph("1.2.3.2. Mẫu hóa đơn 2: Khổ 80mm – Đường viền tách giữa từng dòng sản phẩm")
    doc.add_paragraph("Mẫu này có đường kẻ ngang ngăn cách giữa các sản phẩm.")
    doc.add_paragraph("Ưu điểm:")
    doc.add_paragraph("- Dễ phân biệt từng sản phẩm")
    doc.add_paragraph("- Trông chuyên nghiệp hơn")
    doc.add_paragraph("- Phù hợp với đơn hàng nhiều sản phẩm")

    # 1.2.3.3
    doc.add_paragraph("1.2.3.3. Mẫu hóa đơn 3: Khổ 80mm – Đóng khung từng dòng sản phẩm")
    doc.add_paragraph("Mẫu này đóng khung hoàn toàn cho từng sản phẩm.")
    doc.add_paragraph("Đặc điểm:")
    doc.add_paragraph("- Mỗi sản phẩm nằm trong một ô riêng biệt")
    doc.add_paragraph("- Nổi bật, dễ kiểm tra")
    doc.add_paragraph("- Phù hợp cho cửa hàng cao cấp")

    # 1.2.3.4
    doc.add_paragraph("1.2.3.4. Mẫu hóa đơn 4: Khổ A4/A5")
    doc.add_paragraph("Mẫu hóa đơn khổ lớn cho các giao dịch quan trọng.")
    doc.add_paragraph("Sử dụng khi:")
    doc.add_paragraph("- Đơn hàng giá trị cao")
    doc.add_paragraph("- Cần ghi chú chi tiết")
    doc.add_paragraph("- Hóa đơn VAT")
    doc.add_paragraph("Nội dung bao gồm:")
    doc.add_paragraph("1 Thông tin công ty đầy đủ")
    doc.add_paragraph("2 Thông tin khách hàng")
    doc.add_paragraph("3 Bảng chi tiết sản phẩm")
    doc.add_paragraph("4 Chữ ký người bán và người mua")

    # 2. Quản lý kho hàng
    doc.add_paragraph("2. Quản lý kho hàng")
    doc.add_paragraph("Quản lý tồn kho hiệu quả giúp tránh thiếu hụt hoặc tồn đọng hàng.")

    # 2.1. Nhập hàng
    doc.add_paragraph("2.1. Nhập hàng")
    doc.add_paragraph("Quy trình nhập hàng vào kho:")
    doc.add_paragraph("1 Tạo phiếu nhập hàng")
    doc.add_paragraph("2 Nhập thông tin nhà cung cấp")
    doc.add_paragraph("3 Nhập danh sách sản phẩm")
    doc.add_paragraph("- Mã sản phẩm")
    doc.add_paragraph("- Số lượng")
    doc.add_paragraph("- Giá nhập")
    doc.add_paragraph("4 Lưu phiếu nhập")
    doc.add_paragraph("5 Hệ thống tự động cập nhật tồn kho")

    # 2.2. Xuất hàng
    doc.add_paragraph("2.2. Xuất hàng")
    doc.add_paragraph("Khi có đơn hàng, hệ thống tự động xuất hàng.")
    doc.add_paragraph("Bạn cũng có thể xuất hàng thủ công:")
    doc.add_paragraph("1 Tạo phiếu xuất hàng")
    doc.add_paragraph("2 Chọn lý do xuất (bán hàng, hỏng, khuyến mãi)")
    doc.add_paragraph("3 Nhập sản phẩm và số lượng")
    doc.add_paragraph("4 Xác nhận xuất hàng")

    # Lưu file
    doc.save("tendoo_guide.docx")
    print("Đã tạo file tendoo_guide.docx thành công!")


if __name__ == "__main__":
    create_sample_tendoo_document()
